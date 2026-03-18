"""
Gera o documento acadêmico PIM V em formato ABNT (.docx)
Seções: Capa, Folha de Rosto, Resumo, Abstract, Sumário, Introdução
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================
# Configurações ABNT
# ============================================================
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
FONT_SIZE_SMALL = Pt(10)  # citações longas, notas de rodapé
LINE_SPACING = 1.5
MARGIN_TOP = Cm(3)
MARGIN_BOTTOM = Cm(2)
MARGIN_LEFT = Cm(3)
MARGIN_RIGHT = Cm(2)
FIRST_LINE_INDENT = Cm(1.5)

# Dados do aluno
ALUNO = "ENZO XAVIER SANTOS"
RA = "2620639"
CIDADE = "São Paulo"  # Ajustar se necessário
ANO = "2026"
CURSO = "Análise e Desenvolvimento de Sistemas"
ORIENTADOR = "Prof. MSc. Tarcísio Peres"


def configurar_secao(section):
    """Configura margens ABNT para a seção."""
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT


def set_paragraph_font(paragraph, font_name=FONT_NAME, font_size=FONT_SIZE, bold=False, color=None):
    """Define fonte padrão para um parágrafo."""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        if color:
            run.font.color.rgb = color


def add_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
                  font_size=FONT_SIZE, spacing_after=Pt(0), spacing_before=Pt(0),
                  first_line_indent=None, font_name=FONT_NAME, line_spacing=LINE_SPACING):
    """Adiciona parágrafo formatado ABNT."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = spacing_after
    p.paragraph_format.space_before = spacing_before
    p.paragraph_format.line_spacing = line_spacing

    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent

    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold

    # Garantir fonte no XML (para fontes especiais)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)

    return p


def add_empty_lines(doc, count=1):
    """Adiciona linhas em branco."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = LINE_SPACING
        run = p.add_run("")
        run.font.size = FONT_SIZE
        run.font.name = FONT_NAME


def add_page_break(doc):
    """Adiciona quebra de página."""
    doc.add_page_break()


def set_page_number(section, start_from=None):
    """Configura numeração de página no canto superior direito."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if start_from is not None:
        # Adiciona campo de número de página
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)

        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._element.append(instrText)

        run3 = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._element.append(fldChar2)

        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = FONT_SIZE


# ============================================================
# CAPA
# ============================================================
def criar_capa(doc):
    section = doc.sections[0]
    configurar_secao(section)

    # Sem numeração na capa
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.text = ""

    add_empty_lines(doc, 2)

    # UNIVERSIDADE PAULISTA – UNIP EaD
    add_paragraph(doc, "UNIVERSIDADE PAULISTA – UNIP EaD",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    add_empty_lines(doc, 1)

    # Projeto Integrado Multidisciplinar
    add_paragraph(doc, "Projeto Integrado Multidisciplinar",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Curso Superior de Tecnologia em ADS
    add_paragraph(doc, f"Curso Superior de Tecnologia em {CURSO}",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_empty_lines(doc, 3)

    # Nome do aluno e RA
    add_paragraph(doc, f"{ALUNO} – RA: {RA}",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    add_empty_lines(doc, 4)

    # Título
    add_paragraph(doc, "AMBIENTE VIRTUAL RESPONSIVO PARA EVENTOS",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                  font_size=Pt(14))
    add_paragraph(doc, "ACADÊMICOS INCLUSIVOS EM TI",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                  font_size=Pt(14))

    add_empty_lines(doc, 1)

    # Subtítulo
    add_paragraph(doc, "Semana de TI Inclusiva UNIP 2026",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=Pt(12))

    add_empty_lines(doc, 8)

    # Cidade e ano
    add_paragraph(doc, CIDADE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, ANO, alignment=WD_ALIGN_PARAGRAPH.CENTER)


# ============================================================
# FOLHA DE ROSTO
# ============================================================
def criar_folha_rosto(doc):
    add_page_break(doc)

    add_empty_lines(doc, 4)

    # Nome do aluno
    add_paragraph(doc, f"{ALUNO} – RA: {RA}",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    add_empty_lines(doc, 5)

    # Título
    add_paragraph(doc, "AMBIENTE VIRTUAL RESPONSIVO PARA EVENTOS",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                  font_size=Pt(14))
    add_paragraph(doc, "ACADÊMICOS INCLUSIVOS EM TI",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                  font_size=Pt(14))

    add_empty_lines(doc, 4)

    # Natureza do trabalho (recuado à direita)
    natureza = (
        "Projeto Integrado Multidisciplinar para obtenção do título de tecnólogo "
        f"em {CURSO}, apresentado à Universidade Paulista – UNIP EaD."
    )
    p = add_paragraph(doc, natureza,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      font_size=FONT_SIZE_SMALL,
                      line_spacing=1.0)
    # Recuo à direita (metade da página)
    p.paragraph_format.left_indent = Cm(7)

    add_empty_lines(doc, 1)

    # Orientador
    p_orient = add_paragraph(doc, f"Orientador: {ORIENTADOR}",
                             alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             font_size=FONT_SIZE_SMALL)
    p_orient.paragraph_format.left_indent = Cm(7)

    add_empty_lines(doc, 8)

    # Cidade e ano
    add_paragraph(doc, CIDADE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, ANO, alignment=WD_ALIGN_PARAGRAPH.CENTER)


# ============================================================
# RESUMO
# ============================================================
def criar_resumo(doc):
    add_page_break(doc)

    add_paragraph(doc, "RESUMO",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                  spacing_after=Pt(24))

    texto_resumo = (
        "Este trabalho apresenta o desenvolvimento de um ambiente virtual responsivo "
        "para a gestão de eventos acadêmicos inclusivos em Tecnologia da Informação, "
        "denominado \"Semana de TI Inclusiva UNIP 2026\". O projeto integra quatro "
        "disciplinas do curso de Análise e Desenvolvimento de Sistemas: Desenvolvimento "
        "Web Responsivo, Programação Orientada a Objetos com C#, Comunicação, Liderança "
        "e Negociação, e Língua Brasileira de Sinais (Libras). A aplicação consiste em "
        "um front-end responsivo construído com HTML5, CSS3 e JavaScript vanilla, "
        "seguindo a abordagem mobile-first com breakpoints para smartphones (375px), "
        "tablets (768px) e desktops (1280px), e um back-end desenvolvido em ASP.NET Core "
        "com C#, utilizando arquitetura REST e persistência em arquivos CSV. O sistema "
        "permite o gerenciamento completo do ciclo de vida do evento: cadastro de "
        "participantes, inscrição em atividades, consulta de programação com filtros por "
        "tipo e dia, emissão e validação de certificados digitais, área do participante "
        "com notificações e envio de feedback. Como diferencial de inclusão, o ambiente "
        "integra um glossário de termos técnicos de TI traduzidos para Libras, com "
        "descrição dos sinais correspondentes, categorização temática e indicação de "
        "contexto de uso, além de oferecer a possibilidade de solicitação de intérprete "
        "de Libras em cada atividade. A modelagem orientada a objetos emprega herança, "
        "interfaces, enumerações e coleções genéricas, evidenciando os princípios de "
        "encapsulamento e polimorfismo. O plano de comunicação contempla o mapeamento "
        "de públicos, canais, mensagens-chave e cenários de negociação para a gestão "
        "do evento. Os resultados demonstram que a convergência entre desenvolvimento "
        "web, programação orientada a objetos, comunicação organizacional e Libras "
        "permite a construção de um ecossistema acadêmico mais inclusivo, replicável "
        "em diferentes cursos e instituições."
    )
    add_paragraph(doc, texto_resumo,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_line_indent=FIRST_LINE_INDENT,
                  line_spacing=1.0)

    add_empty_lines(doc, 1)

    add_paragraph(doc, "Palavras-chave: Desenvolvimento Web Responsivo. Programação Orientada a Objetos. "
                  "Acessibilidade Digital. Libras. Eventos Acadêmicos Inclusivos.",
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  bold=True, font_size=FONT_SIZE,
                  line_spacing=1.0)


# ============================================================
# ABSTRACT
# ============================================================
def criar_abstract(doc):
    add_page_break(doc)

    add_paragraph(doc, "ABSTRACT",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                  spacing_after=Pt(24))

    texto_abstract = (
        "This paper presents the development of a responsive virtual environment "
        "for managing inclusive academic events in Information Technology, named "
        "\"Inclusive IT Week UNIP 2026\". The project integrates four disciplines "
        "from the Systems Analysis and Development program: Responsive Web Development, "
        "Object-Oriented Programming with C#, Communication, Leadership and Negotiation, "
        "and Brazilian Sign Language (Libras). The application consists of a responsive "
        "front-end built with HTML5, CSS3 and vanilla JavaScript, following a mobile-first "
        "approach with breakpoints for smartphones (375px), tablets (768px) and desktops "
        "(1280px), and a back-end developed in ASP.NET Core with C#, using REST "
        "architecture and CSV file persistence. The system enables complete event "
        "lifecycle management: participant registration, activity enrollment, schedule "
        "viewing with filters by type and day, digital certificate issuance and "
        "validation, participant dashboard with notifications and feedback submission. "
        "As an inclusion differentiator, the environment integrates a glossary of IT "
        "technical terms translated into Libras, with corresponding sign descriptions, "
        "thematic categorization and usage context, in addition to offering the option "
        "to request a Libras interpreter for each activity. The object-oriented modeling "
        "employs inheritance, interfaces, enumerations and generic collections, "
        "demonstrating the principles of encapsulation and polymorphism. The communication "
        "plan covers audience mapping, channels, key messages and negotiation scenarios "
        "for event management. The results demonstrate that the convergence of web "
        "development, object-oriented programming, organizational communication and "
        "Libras enables the construction of a more inclusive academic ecosystem, "
        "replicable across different programs and institutions."
    )
    add_paragraph(doc, texto_abstract,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_line_indent=FIRST_LINE_INDENT,
                  line_spacing=1.0)

    add_empty_lines(doc, 1)

    add_paragraph(doc, "Keywords: Responsive Web Development. Object-Oriented Programming. "
                  "Digital Accessibility. Brazilian Sign Language. Inclusive Academic Events.",
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  bold=True, font_size=FONT_SIZE,
                  line_spacing=1.0)


# ============================================================
# SUMÁRIO
# ============================================================
def criar_sumario(doc):
    add_page_break(doc)

    add_paragraph(doc, "SUMÁRIO",
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
                  spacing_after=Pt(24))

    itens = [
        ("1", "INTRODUÇÃO", "X"),
        ("2", "DESENVOLVIMENTO WEB RESPONSIVO", "X"),
        ("2.1", "Estrutura e organização do front-end", "X"),
        ("2.2", "Layout responsivo e media queries", "X"),
        ("2.3", "Acessibilidade e boas práticas", "X"),
        ("2.4", "Testes de responsividade", "X"),
        ("3", "PROGRAMAÇÃO ORIENTADA A OBJETOS COM C#", "X"),
        ("3.1", "Modelagem orientada a objetos", "X"),
        ("3.2", "Arquitetura da API REST", "X"),
        ("3.3", "Persistência em arquivos CSV", "X"),
        ("3.4", "Tratamento de exceções e validações", "X"),
        ("4", "COMUNICAÇÃO, LIDERANÇA E NEGOCIAÇÃO", "X"),
        ("4.1", "Mapeamento de públicos e canais", "X"),
        ("4.2", "Mensagens-chave e estratégias de divulgação", "X"),
        ("4.3", "Cenários de negociação e gestão de conflitos", "X"),
        ("5", "LÍNGUA BRASILEIRA DE SINAIS – LIBRAS", "X"),
        ("5.1", "Glossário de termos de TI em Libras", "X"),
        ("5.2", "Roteiros para conteúdo em Libras", "X"),
        ("5.3", "Adequação linguística e perspectiva bilíngue", "X"),
        ("6", "CONCLUSÃO", "X"),
        ("", "REFERÊNCIAS", "X"),
    ]

    for num, titulo, pagina in itens:
        # Itens principais em negrito
        is_main = num in ("1", "2", "3", "4", "5", "6", "")
        prefix = f"{num} " if num else ""
        text = f"{prefix}{titulo}"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(2)

        if not num.startswith("") or num == "":
            # Sub-itens com recuo
            if "." in num:
                p.paragraph_format.left_indent = Cm(1)

        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.font.bold = is_main

        # Pontilhado + número da página (placeholder)
        dots = " " + "." * 60 + " "
        run2 = p.add_run(dots)
        run2.font.name = FONT_NAME
        run2.font.size = Pt(10)

        run3 = p.add_run(pagina)
        run3.font.name = FONT_NAME
        run3.font.size = FONT_SIZE
        run3.font.bold = is_main


# ============================================================
# INTRODUÇÃO
# ============================================================
def criar_introducao(doc, section_number):
    add_page_break(doc)

    # Aqui começa a numeração de página
    add_paragraph(doc, "1 INTRODUÇÃO",
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=True,
                  font_size=FONT_SIZE, spacing_after=Pt(18))

    paragrafos = [
        (
            "A crescente digitalização dos processos acadêmicos tem evidenciado a necessidade "
            "de plataformas virtuais que não apenas gerenciem eventos de forma eficiente, mas "
            "que também garantam a acessibilidade e a inclusão de todos os participantes. No "
            "contexto brasileiro, a Lei nº 10.436/2002, que reconhece a Língua Brasileira de "
            "Sinais (Libras) como meio legal de comunicação das comunidades surdas, e o "
            "Decreto nº 5.626/2005, que regulamenta essa lei, estabelecem diretrizes claras "
            "sobre a necessidade de recursos em Libras em ambientes educacionais. Somam-se a "
            "essas legislações a Lei Brasileira de Inclusão (Lei nº 13.146/2015), que "
            "determina a promoção da acessibilidade em plataformas digitais, e as Diretrizes "
            "de Acessibilidade para Conteúdo Web (WCAG), que orientam o desenvolvimento de "
            "interfaces acessíveis."
        ),
        (
            "Eventos acadêmicos em Tecnologia da Informação, como semanas de curso, "
            "congressos estudantis, competições de programação e mostras de projetos, "
            "constituem espaços privilegiados de circulação de conhecimento, criação de redes "
            "profissionais e fortalecimento da identidade dos cursos. A expansão de formatos "
            "híbridos e remotos intensifica a dependência de plataformas digitais para "
            "inscrição, acompanhamento de programação, interação com palestrantes e acesso a "
            "certificados, ao mesmo tempo em que expõe desigualdades de acesso linguístico e "
            "comunicacional entre diferentes grupos de participantes (SASSAKI, 2009)."
        ),
        (
            "Diante desse cenário, o presente trabalho tem como objetivo o desenvolvimento de "
            "um ambiente virtual responsivo para a gestão de eventos acadêmicos inclusivos em "
            "Tecnologia da Informação, denominado \"Semana de TI Inclusiva UNIP 2026\". A "
            "proposta integra quatro disciplinas do curso de Análise e Desenvolvimento de "
            "Sistemas da Universidade Paulista: Desenvolvimento Web Responsivo, Programação "
            "Orientada a Objetos com C#, Comunicação, Liderança e Negociação, e Língua "
            "Brasileira de Sinais (Libras)."
        ),
        (
            "O ambiente virtual desenvolvido compreende um front-end responsivo construído "
            "com HTML5, CSS3 e JavaScript, adotando a abordagem mobile-first com pontos de "
            "quebra para smartphones, tablets e desktops, e um back-end implementado em "
            "ASP.NET Core com C#, seguindo a arquitetura REST para comunicação entre cliente "
            "e servidor. O sistema permite o gerenciamento completo do ciclo de vida do "
            "evento, incluindo cadastro de participantes, inscrição em atividades, consulta "
            "da programação, emissão de certificados digitais e envio de feedback."
        ),
        (
            "Como diferencial de inclusão, o projeto implementa um glossário de termos "
            "técnicos de Tecnologia da Informação traduzidos para Libras, com descrição "
            "dos sinais correspondentes, categorização temática e indicação de contexto de "
            "uso. Além disso, oferece a possibilidade de solicitação de intérprete de Libras "
            "em cada atividade cadastrada, contribuindo para a participação plena de pessoas "
            "surdas no evento acadêmico."
        ),
        (
            "A metodologia adotada consiste em pesquisa aplicada com abordagem prática, "
            "fundamentada em revisão bibliográfica sobre desenvolvimento web responsivo "
            "(MARCOTTE, 2011), programação orientada a objetos (DEITEL; DEITEL, 2017), "
            "comunicação organizacional (KUNSCH, 2003) e educação bilíngue para surdos "
            "(QUADROS, 2006). O desenvolvimento seguiu o ciclo de análise de requisitos, "
            "modelagem, implementação e testes, utilizando ferramentas como Visual Studio "
            "Code, .NET SDK 9, navegadores para testes de responsividade e o Swagger para "
            "documentação e testes da API."
        ),
        (
            "O presente trabalho está organizado da seguinte forma: o Capítulo 2 aborda o "
            "Desenvolvimento Web Responsivo, descrevendo a estrutura do front-end, as "
            "técnicas de layout responsivo e as práticas de acessibilidade adotadas; o "
            "Capítulo 3 trata da Programação Orientada a Objetos com C#, apresentando a "
            "modelagem das entidades, a arquitetura da API REST e os mecanismos de "
            "persistência; o Capítulo 4 discute a Comunicação, Liderança e Negociação, "
            "com o plano de comunicação do evento e cenários de negociação; o Capítulo 5 "
            "aborda a Língua Brasileira de Sinais, incluindo o glossário de termos técnicos "
            "e as diretrizes de adequação linguística; e, por fim, o Capítulo 6 apresenta "
            "as conclusões do trabalho."
        ),
    ]

    for texto in paragrafos:
        add_paragraph(doc, texto,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      first_line_indent=FIRST_LINE_INDENT,
                      spacing_after=Pt(6))


# ============================================================
# MAIN — Gerar documento
# ============================================================
def main():
    doc = Document()

    # Configurar estilos padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    style.paragraph_format.line_spacing = LINE_SPACING

    # Garantir fonte no XML
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    rPr.insert(0, rFonts)

    # Gerar seções
    criar_capa(doc)
    criar_folha_rosto(doc)
    criar_resumo(doc)
    criar_abstract(doc)
    criar_sumario(doc)
    criar_introducao(doc, 1)

    # Salvar
    output_path = os.path.join(os.path.dirname(__file__), "..", "docs", "PIM_V_Documento.docx")
    output_path = os.path.normpath(output_path)
    doc.save(output_path)
    print(f"Documento gerado com sucesso: {output_path}")
    print(f"Seções: Capa, Folha de Rosto, Resumo, Abstract, Sumário, Introdução")


if __name__ == "__main__":
    main()
